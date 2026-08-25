"""Общие стили Word-документов ООО «ФИТЭРА» для программных контрактов SNAX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run

ROOT = Path(__file__).resolve().parents[1]
LOGO_PATH = ROOT / "docs" / "assets" / "fitera-mark.png"

GREEN = RGBColor(0x1F, 0x6D, 0x45)
GREEN_CONTACT = RGBColor(0x1F, 0x6B, 0x43)
NAVY = RGBColor(0x17, 0x36, 0x5D)
GRAY = RGBColor(0x66, 0x66, 0x66)
DARK = RGBColor(0x33, 0x33, 0x33)
BODY = RGBColor(0x20, 0x21, 0x24)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
YELLOW_TEXT = RGBColor(0xD6, 0x9E, 0x2E)

FILL_GREEN = "EAF4E3"
FILL_BLUE = "EEF5FB"
FILL_BLUE_DEEP = "D9EAF7"
FILL_YELLOW = "FFF2CC"
FILL_HEADER = "1F6D45"
LIME = "8DC63F"
YELLOW_BORDER = "F1B434"

CONTENT_WIDTH_MM = 176


def _rfonts(element, ascii_name: str, east_name: str = "Arial") -> None:
    rPr = element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), ascii_name)
    rFonts.set(qn("w:hAnsi"), ascii_name)
    rFonts.set(qn("w:cs"), east_name)
    rFonts.set(qn("w:eastAsia"), east_name)


def clear_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._p
    for child in list(element):
        if child.tag != qn("w:pPr"):
            element.remove(child)


def set_run_font(
    run: Run,
    *,
    name: str = "Aptos",
    size_pt: float | None = None,
    bold: bool | None = None,
    color: RGBColor | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    _rfonts(run._element, name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_text(
    paragraph: Paragraph,
    text: str,
    *,
    name: str = "Aptos",
    size_pt: float = 11,
    bold: bool = False,
    color: RGBColor = BODY,
    italic: bool = False,
) -> Run:
    run = paragraph.add_run(text)
    set_run_font(run, name=name, size_pt=size_pt, bold=bold, color=color, italic=italic)
    return run


def set_paragraph_spacing(
    paragraph: Paragraph,
    *,
    before: float = 0,
    after: float = 6,
    line: float | None = 1.15,
    align: WD_ALIGN_PARAGRAPH | None = None,
    keep_with_next: bool = False,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    if line is not None:
        fmt.line_spacing = line
        fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if align is not None:
        paragraph.alignment = align
    fmt.keep_with_next = keep_with_next


def shade_cell(cell: _Cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell: _Cell, top: int, start: int, bottom: int, end: int) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcMar"))
    if existing is not None:
        tcPr.remove(existing)
    tc_mar = OxmlElement("w:tcMar")
    for tag, value in (
        ("top", top),
        ("left", start),
        ("bottom", bottom),
        ("right", end),
    ):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tcPr.append(tc_mar)


def set_cell_borders(
    cell: _Cell,
    *,
    color: str,
    size: str = "8",
    sides: tuple[str, ...] = ("top", "left", "bottom", "right"),
    extra: dict[str, tuple[str, str]] | None = None,
) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    borders = OxmlElement("w:tcBorders")
    specified = extra or {}
    for side in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{side}")
        if side in specified:
            val, col = specified[side]
            node.set(qn("w:val"), val)
            node.set(qn("w:sz"), size if val != "nil" else "0")
            node.set(qn("w:color"), col)
        elif side in sides:
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), size)
            node.set(qn("w:space"), "0")
            node.set(qn("w:color"), color)
        else:
            node.set(qn("w:val"), "nil")
        borders.append(node)
    tcPr.append(borders)


def valign_center(cell: _Cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = tcPr.find(qn("w:vAlign"))
    if vAlign is None:
        vAlign = OxmlElement("w:vAlign")
        tcPr.append(vAlign)
    vAlign.set(qn("w:val"), "center")


def configure_table(table: Table, widths_mm: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if hasattr(table, "allow_autofit"):
        table.allow_autofit = False
    widths_dxa = [int(width * 56.7) for width in widths_mm]
    total = sum(widths_dxa)
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    jc = tbl_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        tbl_pr.append(jc)
    jc.set(qn("w:val"), "center")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for index, width in enumerate(widths_dxa):
            grid[index].set(qn("w:w"), str(width))
    for row in table.rows:
        for index, width in enumerate(widths_dxa):
            cell = row.cells[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def prevent_row_split(table: Table) -> None:
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant = tr_pr.find(qn("w:cantSplit"))
        if cant is None:
            cant = OxmlElement("w:cantSplit")
            tr_pr.append(cant)


def clear_cell(cell: _Cell) -> Paragraph:
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    paragraph = cell.paragraphs[0]
    clear_paragraph(paragraph)
    return paragraph


def add_page_number(paragraph: Paragraph) -> None:
    run = paragraph.add_run("Стр. ")
    set_run_font(run, size_pt=9, color=GRAY)
    begin = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    begin._r.append(fld_begin)
    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    instr_run._r.append(instr)
    end = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end._r.append(fld_end)


def add_paragraph_border(paragraph: Paragraph, color: str = LIME, size: str = "12") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY
    _rfonts(normal.element, "Aptos")
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    title = styles["Title"]
    title.font.name = "Aptos"
    title.font.size = Pt(25)
    title.font.bold = True
    title.font.color.rgb = NAVY
    _rfonts(title.element, "Aptos")
    title.paragraph_format.space_after = Pt(6)
    title.paragraph_format.line_spacing = 1.0

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Aptos"
    subtitle.font.size = Pt(13)
    subtitle.font.color.rgb = GREEN
    subtitle.font.italic = False
    _rfonts(subtitle.element, "Aptos")
    subtitle.paragraph_format.space_after = Pt(8)

    heading1 = styles["Heading 1"]
    heading1.font.name = "Aptos"
    heading1.font.size = Pt(17)
    heading1.font.bold = True
    heading1.font.color.rgb = GREEN
    _rfonts(heading1.element, "Aptos")
    heading1.paragraph_format.space_before = Pt(16)
    heading1.paragraph_format.space_after = Pt(8)

    heading2 = styles["Heading 2"]
    heading2.font.name = "Aptos"
    heading2.font.size = Pt(12.5)
    heading2.font.bold = True
    heading2.font.color.rgb = GREEN
    _rfonts(heading2.element, "Aptos")
    heading2.paragraph_format.space_before = Pt(10)
    heading2.paragraph_format.space_after = Pt(4)

    heading3 = styles["Heading 3"]
    heading3.font.name = "Aptos"
    heading3.font.size = Pt(11.5)
    heading3.font.bold = True
    heading3.font.color.rgb = NAVY
    _rfonts(heading3.element, "Aptos")
    heading3.paragraph_format.space_before = Pt(8)
    heading3.paragraph_format.space_after = Pt(3)


def add_callout(
    document: Document,
    title: str,
    body: str,
    *,
    fill: str = FILL_GREEN,
    border: str = FILL_HEADER,
    title_color: RGBColor = GREEN,
    border_size: str = "10",
    left_only: bool = False,
) -> Table:
    table = document.add_table(rows=1, cols=1)
    configure_table(table, [CONTENT_WIDTH_MM])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    if left_only:
        set_cell_borders(
            cell,
            color=border,
            size="22",
            sides=("left",),
            extra={
                "top": ("nil", "auto"),
                "bottom": ("nil", "auto"),
                "right": ("nil", "auto"),
                "left": ("single", border),
            },
        )
    else:
        set_cell_borders(cell, color=border, size=border_size)
    set_cell_margins(cell, 120, 160, 120, 160)
    title_p = clear_cell(cell)
    set_paragraph_spacing(title_p, after=2, line=1.0)
    add_text(title_p, title, size_pt=11, bold=True, color=title_color)
    body_p = cell.add_paragraph()
    set_paragraph_spacing(body_p, after=0, line=1.1)
    add_text(body_p, body, size_pt=10.5, color=BODY)
    prevent_row_split(table)
    spacer = document.add_paragraph()
    set_paragraph_spacing(spacer, after=8, before=0)
    spacer.paragraph_format.space_after = Pt(8)
    return table


def add_kpi_tiles(document: Document, tiles: list[tuple[str, str]]) -> None:
    cols = 4
    rows = (len(tiles) + cols - 1) // cols
    table = document.add_table(rows=rows, cols=cols)
    configure_table(table, [CONTENT_WIDTH_MM / cols] * cols)
    for index in range(rows * cols):
        row, col = divmod(index, cols)
        cell = table.cell(row, col)
        shade_cell(cell, FILL_BLUE)
        set_cell_margins(cell, 100, 80, 100, 80)
        set_cell_borders(cell, color="FFFFFF", size="8")
        paragraph = clear_cell(cell)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(paragraph, after=0, line=1.0)
        if index < len(tiles):
            value, label = tiles[index]
            add_text(paragraph, value, size_pt=14, bold=True, color=NAVY)
            caption = cell.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(caption, after=0, line=1.0)
            add_text(caption, label, size_pt=8.5, color=GRAY)
    prevent_row_split(table)
    spacer = document.add_paragraph()
    set_paragraph_spacing(spacer, after=6)


def add_data_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_mm: list[float],
) -> Table:
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    configure_table(table, widths_mm)
    for index, header in enumerate(headers):
        cell = table.cell(0, index)
        shade_cell(cell, FILL_HEADER)
        set_cell_margins(cell, 90, 90, 90, 90)
        valign_center(cell)
        paragraph = clear_cell(cell)
        set_paragraph_spacing(paragraph, after=0, line=1.0)
        add_text(paragraph, header, size_pt=10, bold=True, color=WHITE)
    for row_index, row in enumerate(rows, start=1):
        fill = "FFFFFF" if row_index % 2 else FILL_GREEN
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            shade_cell(cell, fill)
            set_cell_margins(cell, 80, 90, 80, 90)
            valign_center(cell)
            paragraph = clear_cell(cell)
            set_paragraph_spacing(paragraph, after=0, line=1.05)
            add_text(paragraph, value, size_pt=10, color=BODY)
    prevent_row_split(table)
    spacer = document.add_paragraph()
    set_paragraph_spacing(spacer, after=8)
    return table


def add_heading(document: Document, text: str, level: int = 1) -> Paragraph:
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_body(document: Document, text: str, *, after: float = 6) -> Paragraph:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after=after)
    add_text(paragraph, text, size_pt=11)
    return paragraph


def add_bullet(document: Document, text: str, *, bold_prefix: str | None = None) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after=3, line=1.12)
    fmt = paragraph.paragraph_format
    fmt.left_indent = Mm(6)
    add_text(paragraph, "•  ", size_pt=11, color=GREEN, bold=True)
    if bold_prefix:
        add_text(paragraph, bold_prefix, size_pt=11, bold=True)
        add_text(paragraph, text, size_pt=11)
    else:
        add_text(paragraph, text, size_pt=11)


def add_numbered(
    document: Document,
    number: int,
    text: str,
    *,
    bold_lead: str | None = None,
) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after=3, line=1.12)
    paragraph.paragraph_format.left_indent = Mm(6)
    add_text(paragraph, f"{number}.  ", size_pt=11, bold=True, color=GREEN)
    if bold_lead:
        add_text(paragraph, bold_lead, size_pt=11, bold=True)
        add_text(paragraph, text, size_pt=11)
    else:
        add_text(paragraph, text, size_pt=11)


def add_checkbox(document: Document, label: str, *, bold_lead: str | None = None) -> None:
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after=4, line=1.15)
    paragraph.paragraph_format.left_indent = Mm(4)
    add_text(paragraph, "☐   ", size_pt=12, color=GREEN)
    if bold_lead:
        add_text(paragraph, bold_lead, size_pt=11, bold=True)
        add_text(paragraph, label, size_pt=11)
    else:
        add_text(paragraph, label, size_pt=11)


def setup_section(
    document: Document,
    *,
    running_header: str,
    footer_line: str,
) -> None:
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(17)
    section.right_margin = Mm(17)
    section.top_margin = Mm(16)
    section.bottom_margin = Mm(16)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(8)
    section.different_first_page_header_footer = True

    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""
    running = section.header.paragraphs[0]
    running.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    clear_paragraph(running)
    add_text(
        running,
        running_header,
        name="Aptos",
        size_pt=8.5,
        bold=True,
        color=GREEN,
    )

    def fill_footer(footer) -> None:
        paragraph = footer.paragraphs[0]
        clear_paragraph(paragraph)
        add_text(paragraph, footer_line, size_pt=8, color=GRAY)
        page = footer.add_paragraph()
        page.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_spacing(page, after=0, before=0, line=1.0)
        add_page_number(page)

    fill_footer(section.first_page_footer)
    fill_footer(section.footer)


def add_letterhead(document: Document) -> None:
    table = document.add_table(rows=1, cols=2)
    configure_table(table, [28, CONTENT_WIDTH_MM - 28])
    logo_cell = table.cell(0, 0)
    text_cell = table.cell(0, 1)
    set_cell_borders(logo_cell, color="FFFFFF", sides=())
    set_cell_borders(text_cell, color="FFFFFF", sides=())
    set_cell_margins(logo_cell, 0, 0, 0, 80)
    set_cell_margins(text_cell, 40, 80, 0, 0)
    valign_center(logo_cell)
    valign_center(text_cell)

    logo_p = clear_cell(logo_cell)
    set_paragraph_spacing(logo_p, after=0, line=1.0)
    if LOGO_PATH.is_file():
        run = logo_p.add_run()
        run.add_picture(str(LOGO_PATH), width=Mm(22))
    else:
        add_text(logo_p, "ФИТЭРА", name="Arial", size_pt=14, bold=True, color=GREEN)

    name_p = clear_cell(text_cell)
    set_paragraph_spacing(name_p, after=1, line=1.0)
    add_text(name_p, "ООО «ФИТЭРА»", name="Arial", size_pt=15, bold=True, color=DARK)
    city = text_cell.add_paragraph()
    set_paragraph_spacing(city, after=0, line=1.0)
    add_text(city, "Приморский край, г. Владивосток", name="Arial", size_pt=9, color=GRAY)
    inn = text_cell.add_paragraph()
    set_paragraph_spacing(inn, after=0, line=1.0)
    add_text(inn, "ИНН 2543052510, КПП 254301001", name="Arial", size_pt=8.5, color=GRAY)
    contact = text_cell.add_paragraph()
    set_paragraph_spacing(contact, after=0, line=1.0)
    add_text(
        contact,
        "8-914-078-94-33  •  info@fitera-dv.ru",
        name="Arial",
        size_pt=9,
        bold=True,
        color=GREEN_CONTACT,
    )
    prevent_row_split(table)

    slogan = document.add_paragraph()
    set_paragraph_spacing(slogan, before=4, after=2, line=1.0)
    add_text(
        slogan,
        "практический взгляд на процессы, учёт и автоматизацию",
        name="Arial",
        size_pt=9,
        italic=True,
        color=GRAY,
    )
    rule = document.add_paragraph()
    set_paragraph_spacing(rule, before=0, after=10, line=1.0)
    add_paragraph_border(rule)


def add_architecture_cards(
    document: Document,
    cards: tuple[tuple[str, str, str], ...] | None = None,
    note: str | None = None,
) -> None:
    if cards is None:
        cards = (
            (
                FILL_BLUE_DEEP,
                "Сервис",
                "Разбирает файлы поставщиков: сохраняет оригинал, "
                "не теряет строки, не исполняет макросы Excel.",
            ),
            (
                FILL_GREEN,
                "1С — master",
                "Номенклатура, связи, остатки, расчёт потребности, заказ и проведение поступления.",
            ),
            (
                FILL_YELLOW,
                "Магазин",
                "Фиксирует факт сканером. Проводит документ только уполномоченный сотрудник.",
            ),
        )
    table = document.add_table(rows=1, cols=len(cards))
    configure_table(table, [CONTENT_WIDTH_MM / len(cards)] * len(cards))
    for index, (fill, title, body) in enumerate(cards):
        cell = table.cell(0, index)
        shade_cell(cell, fill)
        set_cell_margins(cell, 120, 120, 120, 120)
        set_cell_borders(cell, color="FFFFFF", size="12")
        paragraph = clear_cell(cell)
        set_paragraph_spacing(paragraph, after=4, line=1.0)
        add_text(paragraph, title, size_pt=11, bold=True, color=NAVY)
        body_p = cell.add_paragraph()
        set_paragraph_spacing(body_p, after=0, line=1.08)
        add_text(body_p, body, size_pt=9.5, color=BODY)
    prevent_row_split(table)
    if note is None:
        note = (
            "Сервис не считает, сколько заказать, и не создаёт заказ поставщику. "
            "Новые магазины получают данные из центральной 1С, а не сырой файл поставщика. "
            "Это схема «Форус» плюс сервис нормализации как помощник центральной УТ."
        )
    if note:
        note_p = document.add_paragraph()
        set_paragraph_spacing(note_p, before=6, after=8)
        add_text(note_p, note, size_pt=10.5)


def add_code_block(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    configure_table(table, [CONTENT_WIDTH_MM])
    cell = table.cell(0, 0)
    shade_cell(cell, FILL_BLUE)
    set_cell_margins(cell, 100, 120, 100, 120)
    set_cell_borders(cell, color="D9EAF7", size="4")
    paragraph = clear_cell(cell)
    set_paragraph_spacing(paragraph, after=0, line=1.08)
    add_text(paragraph, text, name="Consolas", size_pt=8, color=NAVY)
    prevent_row_split(table)
    spacer = document.add_paragraph()
    set_paragraph_spacing(spacer, after=8)


def add_signature_table(
    document: Document,
    roles: list[str] | None = None,
) -> None:
    if roles is None:
        roles = [
            "Спонсор / собственник",
            "Финансовый директор",
            "Руководитель категорийного менеджмента",
            "Руководитель розницы",
            "Операционный директор",
            "Владелец 1С",
            "Координатор / ФИТЭРА",
        ]
    headers = ["Роль", "ФИО", "Подпись", "Дата"]
    widths = [58, 48, 42, 28]
    table = document.add_table(rows=1 + len(roles), cols=4)
    configure_table(table, widths)
    for index, header in enumerate(headers):
        cell = table.cell(0, index)
        shade_cell(cell, FILL_HEADER)
        set_cell_margins(cell, 80, 90, 80, 90)
        valign_center(cell)
        paragraph = clear_cell(cell)
        set_paragraph_spacing(paragraph, after=0, line=1.0)
        add_text(paragraph, header, size_pt=10, bold=True, color=WHITE)
    for row_index, role in enumerate(roles, start=1):
        for col_index in range(4):
            cell = table.cell(row_index, col_index)
            shade_cell(cell, "FFFFFF")
            set_cell_margins(cell, 140, 90, 80, 90)
            valign_center(cell)
            extra = {
                "top": ("nil", "auto"),
                "left": ("nil", "auto"),
                "right": ("nil", "auto"),
                "bottom": ("single", "B0B0B0"),
            }
            set_cell_borders(cell, color="B0B0B0", extra=extra)
            paragraph = clear_cell(cell)
            set_paragraph_spacing(paragraph, after=0, line=1.0)
            if col_index == 0:
                add_text(paragraph, role, size_pt=10, color=BODY)
            else:
                add_text(paragraph, " ", size_pt=10)
    prevent_row_split(table)
