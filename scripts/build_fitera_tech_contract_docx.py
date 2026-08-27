"""Собрать Word детального технического контракта в стиле документов ООО «ФИТЭРА»."""

# ruff: noqa: E501

from __future__ import annotations

import argparse
import sys
from datetime import UTC, datetime
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT / "scripts") not in sys.path:
    sys.path.insert(0, str(ROOT / "scripts"))

from fitera_docx import (  # noqa: E402
    CONTENT_WIDTH_MM,
    FILL_YELLOW,
    GRAY,
    GREEN,
    YELLOW_BORDER,
    YELLOW_TEXT,
    add_architecture_cards,
    add_body,
    add_bullet,
    add_callout,
    add_checkbox,
    add_code_block,
    add_data_table,
    add_heading,
    add_kpi_tiles,
    add_letterhead,
    add_numbered,
    add_paragraph_border,
    add_signature_table,
    add_text,
    configure_styles,
    set_paragraph_spacing,
    setup_section,
)

DEFAULT_OUTPUT = ROOT / "docs" / "ФИТЭРА_SNAX_Технический_контракт_программы_v3.0.docx"


def _title_block(document: Document) -> None:
    kicker = document.add_paragraph()
    set_paragraph_spacing(kicker, after=2, line=1.0, keep_with_next=True)
    add_text(
        kicker,
        "SNAX  ·  ДЕТАЛЬНЫЙ ТЕХНИЧЕСКИЙ КОНТРАКТ ПРОГРАММЫ",
        size_pt=10,
        bold=True,
        color=GREEN,
    )
    title = document.add_paragraph("Технический контракт программы", style="Title")
    title.paragraph_format.keep_with_next = True
    subtitle = document.add_paragraph(
        "Сводный рабочий контракт для архитектора, 1С, backend, QA и Codex · "
        "все потоки, инварианты, контракты, приёмка и календарь",
        style="Subtitle",
    )
    set_paragraph_spacing(subtitle, after=4, line=1.1)
    version = document.add_paragraph()
    set_paragraph_spacing(version, after=10)
    add_text(
        version,
        "Версия 3.0-tech  ·  25 августа 2026 года  ·  Статус: baseline для G0",
        size_pt=10,
        color=GRAY,
    )


def build_document() -> Document:
    document = Document()
    configure_styles(document)
    setup_section(
        document,
        running_header="ФИТЭРА  ·  SNAX  ·  ТЕХНИЧЕСКИЙ КОНТРАКТ",
        footer_line="ООО «ФИТЭРА»  ·  редакция 3.0-tech  ·  25.08.2026  ·  рабочий контракт программы",
    )
    add_letterhead(document)
    _title_block(document)

    add_callout(
        document,
        "Назначение документа",
        "Собрать в один технический контракт всё, что на 25–26.08.2026 уже зафиксировано: "
        "интервью и реестр v1.2, ТЗ v2.0 и дополнение по приёмке v2.1, SPEC v2.0, ADR-001…004, "
        "топологию «Форус», приём выгрузок, календарь v3.0 и срез репозитория. Документ не заменяет "
        "полевые спецификации readers (SPEC.md) и не разрешает угадывать D-01…D-48.",
    )
    add_callout(
        document,
        "Выгрузки баз",
        "Рабочие DT/CF 1С:УТ / Розница / БП на дату контракта не переданы. Появление выгрузки "
        "наполняет каталоги, а не ломает инварианты. Payload в Git запрещён. "
        "Краткая форма для руководителя — редакция 3.0-exec.",
        fill=FILL_YELLOW,
        border=YELLOW_BORDER,
        title_color=YELLOW_TEXT,
        left_only=True,
    )

    add_kpi_tiles(
        document,
        [
            ("185", "требований v1.2"),
            ("113", "требований P0"),
            ("16", "инвариантов"),
            ("4", "ADR"),
            ("11", "JSON Schema"),
            ("G7", "14.05.2027"),
            ("14", "расширений УТ"),
            ("D-48", "открытых решений"),
        ],
    )

    add_heading(document, "0. Карта документов и приоритет")
    add_body(
        document,
        "При расхождении формулировок требований приоритет у реестра Excel v1.2. "
        "Ломающее архитектурное изменение — только новым ADR и одновременной правкой ТЗ/спецификации. "
        "Машинный контракт — JSON Schema и OpenAPI в каталоге contracts/. "
        "Коллизия NFR-*: в программных документах префикс реестра v1.2 пишется REQ-NFR-*; "
        "NFR ТЗ v2.0 остаются NFR-001…016.",
    )
    add_data_table(
        document,
        ["Документ", "Роль"],
        [
            ["Реестр Excel v1.2", "185 требований, риски, решения — канон бизнеса"],
            ["TECH v3.0 (этот документ)", "Сводная техническая сборка программы"],
            ["SPEC.md v2.0", "Readers, raw workbook, профили, package, API импорта"],
            ["TZ addendum v2.1", "Приёмка, receipt-package, рабочее место магазина"],
            ["SPEC_PROGRAM.md v3.0", "Топология, дампы, финансы, BI, расширения"],
            ["ADR-001…004", "Архитектурные решения"],
            ["SCHEDULE.md / DATA_DUMP_INTAKE.md", "Календарь и процедура выгрузок"],
        ],
        [52, CONTENT_WIDTH_MM - 52],
    )

    add_heading(document, "1. Основание")
    add_numbered(
        document,
        1,
        "Интервью: собственник, КМ/склад, финансовый директор, "
        "соучредитель, операционный директор.",
    )
    add_numbered(
        document,
        2,
        "Реестр v1.2 от 26.08.2026: 185 требований (113 P0 / 61 P1 / 11 P2), "
        "34 процесса, 32 риска, 46 решений.",
    )
    add_numbered(document, 3, "ТЗ v2.0, дополнение по приёмке v2.1, SPEC v2.0.")
    add_numbered(
        document,
        4,
        "ADR-001 гибрид; ADR-002 приёмка; ADR-003 outbox/worker; ADR-004 Форус + спутник УТ.",
    )
    add_numbered(document, 5, "Презентация «Форус» 21.08.2026; первичный реестр 14 расширений УТ.")
    add_numbered(
        document,
        6,
        "Срез репозитория 25.08.2026: WORK-001…004 в main; WORK-005…009 вне main.",
    )
    add_callout(
        document,
        "Формула программы",
        "Достоверные данные + утверждённые правила + назначенные владельцы + прозрачные исключения "
        "+ ограниченный пилот = безопасная автоматизация и масштабирование.",
    )

    add_heading(document, "2. Предмет поставки и не-цели")
    add_body(
        document,
        "Наименование: «SNAX — программа стабилизации данных, централизованного заказа, "
        "приёмки поставок и управленческого контура». Исполнитель — ООО «ФИТЭРА».",
    )
    add_numbered(
        document, 1, "Внешний сервис безопасной нормализации прайсов и документов поставки."
    )
    add_numbered(
        document,
        2,
        "Расширение 1С:УТ: staging, matching, apply, РМ приёмки, контролируемый обмен.",
    )
    add_numbered(
        document,
        3,
        "Каталоги MDM / узлов / обменов / расширений и паспорта KPI.",
    )
    add_numbered(
        document,
        4,
        "Пилоты заказа и приёмки, затем отчётность на сверенных источниках.",
    )
    add_numbered(document, 5, "Эксплуатация: мониторинг, backup/restore, обучение, handover.")
    add_callout(
        document,
        "Не-цели первой волны без change request",
        "Замена 1С/кассы/БП; расчёт потребности и заказ во внешнем сервисе; автосоздание номенклатуры; "
        "OCR как факт приёмки или основание проведения; production BI до G1/G2; таможенный брокер и все "
        "франчайзинговые исключения; все поставщики сразу; TASK-036; виртуальный склад; SCL-011; "
        "полная историческая миграция всех баз.",
        fill=FILL_YELLOW,
        border=YELLOW_BORDER,
        title_color=YELLOW_TEXT,
        left_only=True,
    )

    add_heading(document, "3. Неприкосновенные инварианты")
    add_body(
        document,
        "Изменение любого пункта — только новый ADR и одновременная правка ТЗ/спецификации.",
        after=4,
    )
    invariants = [
        "Ни одна строка файла не теряется молча.",
        "Raw-файл и raw-слой неизменяемы; повтор — новый processing run.",
        "1С — источник истины для номенклатуры и подтверждённых связей.",
        "Сервис не рассчитывает потребность и не создаёт заказ поставщику.",
        "Прямой доступ к БД 1С запрещён.",
        "Макросы, external links и формулы Excel не исполняются.",
        "GTIN хранится строкой и не считается глобально уникальным ключом.",
        "Fuzzy matching формирует кандидатов и не утверждает связь.",
        "Обмены и фоновые операции идемпотентны.",
        "Контракты и профили версионируются.",
        "Секреты и коммерческие payload не попадают в Git.",
        "Сервис не проводит поступление и не меняет складской остаток.",
        "Физический факт приёмки принадлежит 1С и не заменяется OCR.",
        "Неизвестный штрихкод сохраняется как исключение.",
        "Заказ, внешний документ, сессия приёмки и документ 1С прослеживаемо связаны.",
        "Сервис — спутник только центральной УТ; в регионы пишет штатный обмен 1С.",
    ]
    for index, item in enumerate(invariants, start=1):
        add_numbered(document, index, item)
    add_body(
        document,
        "Границы модулей: domain без framework/DB/readers; readers возвращают raw, не canonical items; "
        "normalization не открывает файлы и не пишет в БД напрямую; integration_1c только через канонические "
        "контракты; UI не дублирует server-side валидацию; profile DSL без eval и shell. "
        "Деньги и количества — Decimal/NUMERIC и строки JSON. Время — timezone-aware UTC.",
    )

    add_heading(document, "4. Архитектура")
    add_heading(document, "4.1. Корпоративная топология (ADR-004)", level=2)
    add_code_block(
        document,
        "                    ┌─ БП СНЭКС\n"
        "Единая 1С:УТ 11 ────┼─ БП СНЭКС СИБИРЬ\n"
        "        │           └─ Центральная 1С:Розница ── регионы ── франчайзи\n"
        "        ▲\n"
        "        │ import-package / receipt-package / mapping-sync\n"
        "Сервис нормализации (спутник; нет прямого DB 1С)",
    )
    add_bullet(document, "Корпоративный master — единая 1С:УТ 11 УК СНЭКС.")
    add_bullet(
        document,
        "БП получают контролируемую выгрузку из УТ; магазин автономен в локальной Рознице.",
    )
    add_bullet(
        document,
        "Утверждённые номенклатура и цены идут в Розницу штатным обменом, не сырым файлом поставщика.",
    )
    add_bullet(
        document, "BI читает сверенные витрины и не чинит НСИ. Frontol сосуществует до D-34."
    )
    add_bullet(
        document,
        "14 расширений: disposition keep / refactor / replace / retire (D-37). "
        "ПомощникЗакупок не переносится as-is (CFG-005, R-26).",
    )
    add_body(
        document, "Статусы компонентов до выгрузок: PLANNED | UNVERIFIED | VERIFIED | RETIRED."
    )

    add_heading(document, "4.2. Гибрид и очередь (ADR-001, ADR-003)", level=2)
    add_architecture_cards(
        document,
        note=(
            "Сервис владеет оригиналом, SHA-256, readers, профилями, raw-слоем, DQ и каноническим пакетом. "
            "1С владеет GUID, окончательными связями, расчётом потребности, заказом и проведением. "
            "PostgreSQL — истина очереди: ProcessingRun и OutboxMessage в одной транзакции; "
            "Celery JSON-only at-least-once, очередь snax.import.processing.v1."
        ),
    )

    add_heading(document, "4.3. Стек", level=2)
    add_data_table(
        document,
        ["Слой", "Технология"],
        [
            ["Runtime / API", "Python 3.12, FastAPI, Pydantic 2"],
            ["Данные", "PostgreSQL 16+, SQLAlchemy 2, Alembic, NUMERIC не float"],
            ["Очередь", "Celery + Redis; outbox в PostgreSQL"],
            ["Файлы", "MinIO/S3; openpyxl read_only/data_only; XLS в изолированном worker"],
            ["UI", "React + TypeScript; без server-side валидации в браузере"],
            ["1С", "Расширение УТ без снятия типовой с поддержки"],
        ],
        [42, CONTENT_WIDTH_MM - 42],
    )

    add_heading(document, "5. Потоки ценности")
    add_data_table(
        document,
        ["Поток", "Результат первой волны"],
        [
            ["W0 Управление", "Backlog, роли, gates, decision log"],
            ["W-DI Выгрузки", "Каталоги узлов/объектов/расширений/обменов"],
            ["W1 Стабилизация", "Сверка контрольного магазина/дня"],
            ["W2 Методология", "Паспорта НСИ, заказа, цены, P&L, факта продажи"],
            ["W3 Платформа", "Readers, профили, DQ, package, API"],
            ["W4 Заказ", "Черновик заказа с полной трассировкой"],
            ["W5 Приёмка", "Scan fact и типовой документ без повторного ввода"],
            ["W6 Финансы и BI", "P&L/продажи/остатки с признаком полноты"],
            ["W7 Rollout", "Волны, обучение, SLA, handover"],
        ],
        [42, CONTENT_WIDTH_MM - 42],
    )

    add_heading(document, "6. Конвейер файла поставщика")
    add_numbered(
        document,
        1,
        "POST /imports: multipart, Idempotency-Key, object storage, SHA-256, MIME/signature, лимиты.",
    )
    add_numbered(
        document,
        2,
        "Object key формирует сервис. Внутренний путь хранилища пользователю не возвращается.",
    )
    add_numbered(
        document,
        3,
        "Reader → raw-workbook: 1-based индексы, formulaText + cached value без исполнения макросов и формул.",
    )
    add_numbered(
        document,
        4,
        "Профиль — декларативный YAML. Дрейф шапки → PROFILE_REVIEW; старый профиль молча не применяется.",
    )
    add_numbered(
        document,
        5,
        "Классификация PRODUCT/CATEGORY/HEADER/TOTAL/COMMENT/BLANK/DERIVED/ERROR; нормализация с audit.",
    )
    add_numbered(
        document,
        6,
        "import-package с digest/summary/rows/issues. Дальше 1С: staging → matching → apply → черновик заказа.",
    )
    add_body(
        document,
        "stableRowId = source_file_sha256 + sheet_name + row_number + selected_raw_cell_hash. "
        "Товарный ключ — externalVariantKey в контексте поставщика, не глобально.",
    )

    add_heading(document, "7. Сопоставление")
    add_data_table(
        document,
        ["Ступень", "Правило"],
        [
            ["1. Сохранённая связь", "MATCHED_SAVED; смена критичных полей → review"],
            [
                "2. Код поставщика + упаковка",
                "Только в контексте поставщика; глобальный поиск артикула запрещён",
            ],
            ["3. GTIN + упаковка", "MATCHED_EXACT при одном кандидате и совместимой упаковке"],
            ["4. Код без GTIN", "Первая связь требует подтверждения"],
            ["5. Fuzzy-ранжирование", "Score сохраняется; порог не делает APPROVED"],
            ["6. Нет кандидата", "NEW_ITEM / заявка / DO_NOT_BUY"],
        ],
        [48, CONTENT_WIDTH_MM - 48],
    )
    add_body(
        document,
        "Неизвестный коэффициент упаковки — PACKAGING_CONFLICT. При конфликте версий mapping 1С побеждает.",
    )

    add_heading(document, "8. Приёмка")
    add_body(
        document,
        "TO-BE: заказ в 1С → опционально receipt-package → связь с заказом → магазин сканирует "
        "в ограниченном РМ → неизвестное в exception → товаровед проводит типовой документ. "
        "OCR — подсказка с confidence, не факт и не проведение.",
    )
    add_data_table(
        document,
        ["Маршрут", "Смысл"],
        [
            ["DIRECT_TO_STORE", "Поступление на склад/магазин заказа"],
            ["CENTRAL_WAREHOUSE", "Поступление на центральный склад"],
            ["RECEIPT_AND_TRANSFER", "Поступление + перемещение без повторного ввода"],
            ["CROSS_DOCK", "Вне первой волны без отдельного решения"],
        ],
        [48, CONTENT_WIDTH_MM - 48],
    )
    add_body(
        document,
        "Пакет: RECEIVED → PROFILED → PARSED → PUBLISHED → IMPORTED_1C → LINKED_TO_ORDER → "
        "IN_ACCEPTANCE → SUBMITTED → APPROVED → POSTED. Строка: EXPECTED, MATCHED, SCANNED, "
        "SHORTAGE, OVERAGE, UNPLANNED_ITEM, MAPPING_REQUIRED, MULTIPLE_MATCHES, PACKAGING_CONFLICT, "
        "PRICE_MISMATCH, DAMAGED, REJECTED, ACCEPTED. Повтор SHA-256 того же поставщика — DUPLICATE_DOCUMENT. "
        "Магазин не проводит закупку.",
    )

    add_heading(document, "9. Контракты данных и API")
    add_body(
        document,
        "Breaking change = новая major-версия + ADR. Схемы каталогов программы: 1.0.0. "
        "Денежные и количественные поля — строки decimal. GTIN и идентификаторы — строки. Время — UTC.",
    )
    add_data_table(
        document,
        ["Контракт", "Назначение"],
        [
            ["import-package", "Канонический прайс для staging УТ"],
            ["mapping-sync", "Дельта подтверждённых связей из 1С"],
            ["receipt-package", "Документ поставки"],
            ["raw-workbook", "Результат reader без исполнения формул"],
            ["processing-job-message", "Сообщение очереди; без object payload"],
            ["config-dump-manifest", "Реестр выгрузок; gitPolicy=DO_NOT_COMMIT_PAYLOAD"],
            ["mdm-object-catalog", "Магазин–склад–касса–организация–узел"],
            ["extension-passport", "Паспорт расширения 1С"],
            ["exchange-catalog", "Маршруты БП–УТ–Розница–КСО"],
            ["store-day-reconciliation", "Протокол контрольного магазина/дня (G1)"],
            ["kpi-passport", "Формула показателя; сервис не считает KPI"],
        ],
        [52, CONTENT_WIDTH_MM - 52],
    )
    add_body(
        document,
        "Base path /api/v1. Idempotency-Key на POST. X-Correlation-ID. Ошибка: code, message, details, "
        "field, retryable, correlationId. Оператор: POST/GET /imports, rows, issues, profile, publish. "
        "1С: packages/next, chunks, ACK (ACCEPTED|PARTIAL|REJECTED|RETRY_LATER), mappings:sync. "
        "ACK и mapping delta — только после сохранения транзакции 1С. Dump-upload в public OpenAPI до G0 нет.",
    )

    add_heading(document, "10. Выгрузки информационных баз")
    add_body(
        document,
        "Носители: cf/cfe/XML конфигурации, dt тестовой копии, реестр расширений, обезличенные логи обмена. "
        "Не выгрузка: прайс, УПД, .env. Передача out-of-band + manifest.json. В Git — только checksums.",
    )
    add_data_table(
        document,
        ["Код базы", "Зачем"],
        [
            ["UT_CENTRAL", "Master НСИ, закупки, себестоимость, расширения"],
            ["UT_LEGACY / UT_IMPORT", "История, ДоработкаИмпорта, ПомощникЗакупок"],
            ["RETAIL_CENTRAL", "Цены, права, узлы"],
            ["RETAIL_REGION_*", "Карта магазинов, автономность"],
            ["BP_SNEX / BP_SIBIR", "Обмен, аналитики, непроведённые документы"],
        ],
        [48, CONTENT_WIDTH_MM - 48],
    )
    add_body(
        document,
        "Состояния: RECEIVED → QUARANTINE_SCAN → REGISTERED → PARSED_CATALOGS → VERIFIED → ARCHIVED. "
        "Альтернативы: REJECTED_POLICY, BLOCKED_SECRETS, PARTIAL. Нельзя отключать расширения до инвентаризации. "
        "Acceptance DI-AC-001…006: schema, идемпотентность SHA-256, чистое дерево Git, статус каждой базы, "
        "сверка 14 расширений, нет секретов в манифесте.",
    )

    add_heading(document, "11. Доменные правила реестра v1.2")
    add_bullet(
        document,
        " единый backlog, decision log, среды DEV/TEST/UAT/PROD, backup до массовых операций 1С.",
        bold_prefix="GOV/CFG:",
    )
    add_bullet(
        document,
        " таблица соответствий магазин—склад—касса—организация—узел; до дампа UNVERIFIED; упаковки — D-03.",
        bold_prefix="MDM:",
    )
    add_bullet(
        document,
        " документ либо проведён, либо в exception queue; запись без проведения не есть успех.",
        bold_prefix="INT:",
    )
    add_bullet(
        document,
        " заказ создаёт КМ; расчёт — типовой механизм УТ; сервис заканчивается на import-package.",
        bold_prefix="ORD:",
    )
    add_bullet(
        document,
        " не в сервисе; запрет технических закрытий отрицательных остатков; P&L по начислению; казначейство — D-24.",
        bold_prefix="FIN:",
    )
    add_bullet(
        document,
        " проведённый документ без отгрузки/чека не выручка (OPS-002, D-44).",
        bold_prefix="OPS:",
    )
    add_bullet(
        document,
        " неполный срез не выдаётся за полный (R-30); новая точка не стартует до G1/G2 (R-29).",
        bold_prefix="BI/SCL:",
    )

    add_heading(document, "12. Расширение 1С")
    add_body(
        document,
        "Код только в onec-extension/. Типовая конфигурация не снимается с поддержки. "
        "Запись в типовые объекты — после staging и подтверждения. Остаток поставщика хранится отдельно "
        "от складского. Повтор package ID не создаёт второй документ. Критично: "
        "пр_НоменклатураРасширенныйПремиум и ПомощникЗакупок (закрытый код, собственные данные).",
    )

    add_heading(document, "13. Состояния импорта")
    add_data_table(
        document,
        ["Контур", "Машина состояний"],
        [
            [
                "Сервис",
                "RECEIVED → PROCESSING → VALIDATED/BLOCKED/PROFILE_REVIEW → READY_FOR_1C → "
                "DELIVERED → ACCEPTED_BY_1C | REJECTED_BY_1C → ARCHIVED",
            ],
            [
                "1С staging",
                "NEW → LOADED → MATCHING → READY_TO_APPLY → APPLIED | PARTIAL | REJECTED | CANCELLED | ERROR",
            ],
        ],
        [32, CONTENT_WIDTH_MM - 32],
    )

    add_heading(document, "14. Нефункциональные требования")
    add_body(
        document,
        "ТЗ v2.0 NFR-001…016: прослеживаемость; нет молчаливых потерь; идемпотентность; "
        "p95 5 000 строк ≤ 180 с; ≥ 100 000 строк/сутки; UI p95 ≤ 2 с; доступность 99,5% после MVP; "
        "безопасность файлов; TLS; аудит; observability; backup; среды; профили без релиза ядра; "
        "UI на русском, коды ошибок стабильные английские. Приёмка NFR-RCV: отклик скана ≤ 1 с локально / "
        "≤ 3 с удалённо; сканы не теряются; ≥ 2 000 строк поставки; нет полного payload в логах.",
    )

    add_heading(document, "15. Тесты и статус репозитория")
    add_body(
        document,
        "Обязательны unit/property, golden reader/profile, contract JSON Schema/OpenAPI, transactional DB, "
        "API, UI e2e, сценарии 1С. Golden snapshot не обновляется автоматически. "
        "Программа: dump-manifest, float amount сверки отклоняется, empty MDM code отклоняется, "
        "gitPolicy commit-payload отклоняется.",
    )
    add_data_table(
        document,
        ["Блок", "Статус на 25.08.2026"],
        [
            ["Регистрация файла, SHA-256, MinIO", "MERGED (WORK-002)"],
            ["Очередь / outbox", "MERGED (WORK-003)"],
            ["Raw workbook protocol", "MERGED (WORK-004)"],
            ["XLSX reader / profiles", "OPEN / stacked PR"],
            ["CSV / XLS / DQ / package", "не в main"],
            ["Пилотные профили", "блокировано D-02"],
            ["Staging 1С / заказ / приёмка", "не начато (после G2/G3)"],
            ["GitHub Actions", "не стартует (D-29, billing)"],
        ],
        [58, CONTENT_WIDTH_MM - 58],
    )
    add_body(
        document,
        "Каркас не ускоряет G4: без дампов и утверждённых формул пилот заказа запрещён.",
    )

    add_heading(document, "16. Задачи Codex программы")
    add_data_table(
        document,
        ["Task", "Результат"],
        [
            ["TASK-000…036", "Сохраняются: импорт, matching, apply, приёмка, OCR вне волны"],
            ["TASK-037", "Манифест выгрузки, запрет payload в Git"],
            ["TASK-038", "MDM/узловой каталог"],
            ["TASK-039", "Паспорта расширений и disposition"],
            ["TASK-040", "Каталог обменов и store-day reconciliation"],
            ["TASK-041", "Паспорт KPI без расчёта в сервисе"],
        ],
        [38, CONTENT_WIDTH_MM - 38],
    )

    add_heading(document, "17. Ворота приёмки")
    add_data_table(
        document,
        ["Gate", "Дата A", "Критерий"],
        [
            ["T0", "01.09.2026", "Роли, единый backlog"],
            ["G0", "11.09.2026", "Активы, ADR-004, план выгрузок, D-28"],
            ["G1", "09.10.2026*", "Контрольный магазин/день либо owner расхождения"],
            ["G2", "06.11.2026", "Формулы, НСИ, факт продажи, disposition расширений"],
            ["G3", "22.01.2027", "Readers/DQ/API, приём пакета центральной УТ"],
            ["G4", "19.02.2027", "3–5 поставщиков, 30–50 SKU, черновик без дубля"],
            ["G5", "02.04.2027", "Scan fact, типовой документ, UAT приёмки"],
            ["G6", "23.04.2027", "P&L/продажи/остатки, признак полноты"],
            ["G7", "14.05.2027", "Обучение, SLA, restore drill, handover"],
        ],
        [22, 28, CONTENT_WIDTH_MM - 50],
    )
    add_body(
        document,
        "*G1 не раньше чем манифест + 4 недели, если дампы позже 11.09.2026. "
        "UAT — на конкретном магазине, дне, поставщике и SKU. «Сделали форму» без сверки не приёмка.",
    )

    add_heading(document, "18. Календарь и сценарии")
    add_body(
        document,
        "Содержание P0–P7 сохранено (34 рабочие недели). Добавлены НГ 2027, майские и зависимость G1 от дампов. "
        "Передача сценария A: 14.05.2027. Сценарий B: dumpLagWeeks = ceil((дата_манифеста − 11.09.2026) / 7); "
        "G3 сдвигается, если лаг больше 3 недель. Сценарий C (нет 1С параллельно backend): +6…8 недель, "
        "ориентир конец июня 2027 после D-28. Критический путь: выгрузки → каталоги → G1 → G2 → G4 → G5 → G7; "
        "параллельно P3 на synthetic fixtures → G3.",
    )

    add_heading(document, "19. Открытые решения — не кодировать")
    add_callout(
        document,
        "Запрет угадывания",
        "D-01 релиз УТ; D-02 пилотные поставщики; D-03 упаковки; D-04 цена/пороги; D-05 частичное применение; "
        "D-06 автосопоставление; D-07 создание номенклатуры; D-08 хранение/ПДн; D-09 аутентификация; "
        "D-10 тестовая копия и SKU; D-11… маршруты и документ приёмки; D-20 формула заказа; D-24 казначейство; "
        "D-25 BI; D-28 ресурсы; D-29 CI; D-30 merge chain; D-31…D-40 топология и расширения; D-44 факт продажи; "
        "D-47 состав выгрузок; D-48 календарь. Допустимо: пустые каталоги, JSON Schema, синтетические fixtures.",
        fill=FILL_YELLOW,
        border=YELLOW_BORDER,
        title_color=YELLOW_TEXT,
        left_only=True,
    )

    add_heading(document, "20. Трассировка требований")
    add_data_table(
        document,
        ["Контур", "ID"],
        [
            ["Управление", "GOV-001…012, CFG-001…010"],
            ["НСИ", "MDM-001…012, ARC-005"],
            ["Обмены и топология", "INT-001…010, ARC-001…004, ARC-010"],
            ["Прайсы", "IMP-001…014"],
            ["Заказ / цены / остатки", "ORD-001…012, PRC-001…010, INV-001…010"],
            ["Приёмка", "RCV-001…014"],
            ["Финансы", "FIN-001…018"],
            ["Продажи / BI / масштаб", "SAL, BI, OPS, SCL"],
            ["NFR реестра", "REQ-NFR-001…013"],
        ],
        [48, CONTENT_WIDTH_MM - 48],
    )
    add_body(
        document,
        "Критические риски 20+: R-01…R-06, R-14, R-21, R-23…R-27, R-29…R-31. "
        "Полный индекс — docs/research/2026-08-26/requirements-index.md; канон формулировок — Excel v1.2.",
    )

    add_heading(document, "21. Изменение контракта")
    add_body(
        document,
        "Обратимые дополнения — та же major. Breaking — новая версия schema/endpoint и ADR. "
        "Появление выгрузки не breaking, если заполняются optional поля каталогов 1.x. "
        "Change request: причина → влияние на процесс/данные/контракты/1С/сроки → владелец acceptance → "
        "спонсор для baseline → обновление backlog до разработки.",
    )

    document.add_page_break()
    add_heading(document, "22. Лист согласования")
    add_body(document, "Отметьте итог и заполните роли:", after=4)
    add_checkbox(document, " TECH v3.0 как технический baseline программы.", bold_lead="Утверждаю")
    add_checkbox(document, " с замечаниями.", bold_lead="Утверждаю")
    add_checkbox(document, " — вернуть на доработку.", bold_lead="Не утверждаю")
    add_heading(document, "Обязательные инварианты", level=2)
    add_checkbox(document, "Не строим второй расчёт заказа вне 1С.")
    add_checkbox(document, "Нет прямого доступа сервиса к БД 1С.")
    add_checkbox(document, "OCR не заменяет физический факт приёмки.")
    add_checkbox(document, "Payload выгрузок баз в Git не попадает.")
    add_checkbox(document, "Fuzzy matching не утверждает связь.")

    remarks = document.add_paragraph()
    set_paragraph_spacing(remarks, before=10, after=4)
    add_text(remarks, "Замечания", size_pt=12.5, bold=True, color=GREEN)
    for _ in range(3):
        line = document.add_paragraph()
        set_paragraph_spacing(line, after=10, before=4)
        add_paragraph_border(line, color="B0B0B0", size="6")
        add_text(line, " ", size_pt=11)

    add_signature_table(
        document,
        [
            "Спонсор / собственник",
            "Владелец 1С",
            "Технический лидер сервиса",
            "Финансовый директор",
            "Руководитель категорийного менеджмента",
            "Руководитель розницы",
            "Операционный директор",
            "Координатор / ФИТЭРА",
        ],
    )
    closing = document.add_paragraph()
    set_paragraph_spacing(closing, before=12, after=0)
    add_text(
        closing,
        "После утверждения этот документ становится техническим входом G0 вместе с краткой формой "
        "для руководителя. Детали readers остаются в SPEC.md; детали приёмки — в TZ addendum v2.1.",
        size_pt=10,
        italic=True,
        color=GRAY,
    )

    core = document.core_properties
    core.author = "ООО «ФИТЭРА»"
    core.last_modified_by = "ООО «ФИТЭРА»"
    core.title = "SNAX — технический контракт программы"
    core.subject = "Сводный технический контракт v3.0-tech"
    core.category = "Технический контракт"
    core.comments = (
        "Редакция 3.0-tech. Сводка интервью, реестра v1.2, ТЗ, SPEC, ADR и среза репозитория. "
        "Коммерческие выгрузки баз в файл не входят."
    )
    core.created = datetime(2026, 8, 25, 9, 0, tzinfo=UTC)
    core.modified = datetime(2026, 8, 25, 9, 0, tzinfo=UTC)
    core.revision = 1
    return document


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Собрать Word технического контракта в стиле ФИТЭРА."
    )
    parser.add_argument("-o", "--output", type=Path, default=DEFAULT_OUTPUT, help="Путь к DOCX")
    args = parser.parse_args()
    output: Path = args.output
    output.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(output)
    print(f"wrote {output} ({output.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
