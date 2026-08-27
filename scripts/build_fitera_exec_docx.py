"""Собрать Word краткого программного контракта в стиле документов ООО «ФИТЭРА»."""

from __future__ import annotations

import argparse
import sys
from datetime import UTC, datetime
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT / "scripts") not in sys.path:
    sys.path.insert(0, str(ROOT / "scripts"))

from fitera_docx import (  # noqa: E402
    CONTENT_WIDTH_MM,
    FILL_YELLOW,
    GRAY,
    GREEN,
    YELLOW_BORDER,
    YELLOW_TEXT,
    add_architecture_cards,
    add_body,
    add_callout,
    add_checkbox,
    add_data_table,
    add_heading,
    add_kpi_tiles,
    add_letterhead,
    add_numbered,
    add_paragraph_border,
    add_signature_table,
    add_text,
    configure_styles,
    set_paragraph_spacing,
    setup_section,
)

DEFAULT_OUTPUT = ROOT / "docs" / "ФИТЭРА_SNAX_Программный_контракт_для_руководителя_v3.0-exec.docx"


def build_document() -> Document:
    document = Document()
    configure_styles(document)
    setup_section(
        document,
        running_header="ФИТЭРА  ·  SNAX  ·  ПРОГРАММНЫЙ КОНТРАКТ",
        footer_line="ООО «ФИТЭРА»  ·  редакция 3.0-exec  ·  25.08.2026  ·  для подтверждения G0",
    )
    add_letterhead(document)

    kicker = document.add_paragraph()
    set_paragraph_spacing(kicker, after=2, line=1.0, keep_with_next=True)
    add_text(
        kicker,
        "SNAX  ·  ПРОГРАММА СТАБИЛИЗАЦИИ ДАННЫХ, ЗАКУПОК, ПРИЁМКИ И УПРАВЛЕНЧЕСКОГО КОНТУРА",
        size_pt=10,
        bold=True,
        color=GREEN,
    )

    title = document.add_paragraph("Программный контракт для руководителя", style="Title")
    title.paragraph_format.keep_with_next = True
    subtitle = document.add_paragraph(
        "Краткая форма на подтверждение ворот G0 · границы первой волны и календарь до 14 мая 2027",
        style="Subtitle",
    )
    set_paragraph_spacing(subtitle, after=4, line=1.1)

    version = document.add_paragraph()
    set_paragraph_spacing(version, after=10)
    add_text(
        version,
        "Версия 3.0-exec  ·  25 августа 2026 года  ·  "
        "Статус: проект для управленческого подтверждения",
        size_pt=10,
        color=GRAY,
    )

    add_callout(
        document,
        "Назначение документа",
        "Дать собственнику и руководителям одну страницу решений: что получит бизнес, "
        "как устроена система без жаргона, какие сроки подтверждаем и что нужно от заказчика. "
        "После подписи этот лист становится входом ворот G0. Детальный технический контракт — "
        "PROJECT_CONTRACT_TECH.md. Календарь — SCHEDULE.md.",
    )

    add_kpi_tiles(
        document,
        [
            ("01.09.2026", "старт T0"),
            ("11.09.2026", "ворота G0"),
            ("19.02.2027", "пилот заказа"),
            ("14.05.2027", "передача G7"),
            ("185", "требований v1.2"),
            ("113", "требований P0"),
            ("34 нед.", "горизонт работы"),
            ("3–5", "пилотных поставщиков"),
        ],
    )

    add_heading(document, "1. В одном абзаце")
    add_body(
        document,
        "SNAX сейчас держится на людях, Excel и нескольких базах 1С. Отчёты могут выглядеть "
        "полными, даже если точка не выгрузилась, документ записан, но не проведён, а заказ "
        "собран вручную. Программа сначала восстанавливает достоверность данных и правила, "
        "затем запускает два пилота (заказ поставщику и приёмка), и только после этого — "
        "управленческие отчёты и тиражирование регионов.",
    )
    add_callout(
        document,
        "Главный принцип",
        "Данные и правила → пилот → отчёты и масштабирование. Не наоборот.",
    )

    add_heading(document, "2. Что получит бизнес")
    add_data_table(
        document,
        ["Когда", "Что можно проверить руками"],
        [
            [
                "Октябрь 2026",
                "Один магазин, один день: документы, суммы и остатки сходятся "
                "либо у расхождения есть хозяин.",
            ],
            [
                "Ноябрь 2026",
                "Утверждены формулы заказа, остатка, цены, факта продажи и P&L.",
            ],
            [
                "Февраль 2027",
                "3–5 поставщиков: файл → связанный товар → черновик заказа в 1С без дубля.",
            ],
            [
                "Апрель 2027",
                "Магазин сканирует поставку; товаровед проводит готовый документ, "
                "не набирая строки заново.",
            ],
            [
                "14 мая 2027",
                "Передача: обучение, регламент поддержки, резервное копирование.",
            ],
        ],
        [38, CONTENT_WIDTH_MM - 38],
    )
    add_callout(
        document,
        "Не обещаем в этой волне",
        "Новую 1С «с нуля»; заказ, который считает внешний сайт; автоматическое создание "
        "карточек товара; «умный» разбор фото накладной как факт приёмки; красивый BI до "
        "сверки источников; подключение всех поставщиков сразу.",
        fill=FILL_YELLOW,
        border=YELLOW_BORDER,
        title_color=YELLOW_TEXT,
        left_only=True,
        border_size="22",
    )

    add_heading(document, "3. Как устроена система")
    add_architecture_cards(document)

    add_heading(document, "4. Сроки, которые подтверждаем")
    add_body(
        document,
        "Старт: 1 сентября 2026. Передача: 14 мая 2027. Это 34 недели работы плюс праздники "
        "(Новый год и май). Если выгрузки баз 1С придут позже 11 сентября, контрольный день "
        "сдвигается: дата выгрузки + 4 недели, дальше сдвигается вся цепочка.",
    )
    add_data_table(
        document,
        ["Ворота", "Дата", "Смысл для руководителя"],
        [
            ["Старт", "01.09.2026", "Назначены роли, единый список задач"],
            ["G0", "11.09.2026", "Подтверждён этот лист, доступы, план выгрузок"],
            ["G1", "09.10.2026", "Сверка одного дня"],
            ["G2", "06.11.2026", "Подписаны формулы"],
            ["G3", "22.01.2027", "Платформа разбора файлов готова к пилоту"],
            ["G4", "19.02.2027", "Пилот заказа"],
            ["G5", "02.04.2027", "Пилот приёмки"],
            ["G6", "23.04.2027", "P&L и продажи сходятся до документа"],
            ["G7", "14.05.2027", "Сдано в эксплуатацию"],
        ],
        [28, 32, CONTENT_WIDTH_MM - 60],
    )
    add_body(
        document,
        "Каркас сервиса в репозитории уже начат. Это не ускоряет пилот заказа: без выгрузок 1С "
        "и утверждённых формул пилот запускать нельзя.",
        after=8,
    )

    add_heading(document, "5. Что нужно от заказчика")
    add_numbered(
        document,
        1,
        " спонсора и владельцев: закупки, финансы, розница, продажи, 1С.",
        bold_lead="Назначить",
    )
    add_numbered(
        document,
        2,
        " у прежнего подрядчика исходники, расширения, доступы и резервные копии.",
        bold_lead="Забрать",
    )
    add_numbered(
        document,
        3,
        " выгрузки баз (или согласованный доступ только на чтение). Файлы баз в Git не кладутся.",
        bold_lead="Передать",
    )
    add_numbered(
        document,
        4,
        " 3–5 пилотных поставщиков, контрольный магазин, день и 30–50 товаров.",
        bold_lead="Выбрать",
    )
    add_numbered(
        document,
        5,
        " спорные правила за 5 рабочих дней (цена, упаковка, факт продажи, P&L). "
        "Если решения нет — двигаем срок, а не «пусть разработчик угадает».",
        bold_lead="Решать",
    )
    add_callout(
        document,
        "Бюджет людей (D-28) в этом листе не утверждён",
        "Если команда меньше, чем в паспорте проекта, дата 14.05.2027 пересматривается явно, "
        "а не формулировкой «успеем теми же людьми».",
        fill=FILL_YELLOW,
        border=YELLOW_BORDER,
        title_color=YELLOW_TEXT,
        left_only=True,
    )

    document.add_page_break()
    add_heading(document, "6. Подтверждение руководителя")
    add_body(document, "Отметьте одно:", after=4)
    add_checkbox(
        document,
        " программу, границы первой волны и календарь до 14.05.2027 (при выгрузках до 11.09.2026).",
        bold_lead="Подтверждаю",
    )
    add_checkbox(
        document,
        " (вписать ниже).",
        bold_lead="Подтверждаю с замечаниями",
    )
    add_checkbox(
        document,
        " — вернуть на доработку.",
        bold_lead="Не подтверждаю",
    )

    add_heading(document, "Обязательные галочки, если выбрано «подтверждаю»", level=2)
    add_checkbox(document, "Не строим второй расчёт заказа вне 1С.")
    add_checkbox(
        document,
        "Не публикуем управленческий BI как официальный, "
        "пока не пройдены сверка дня (G1) и формулы (G2).",
    )
    add_checkbox(
        document,
        "Выгрузки баз передаём по отдельному регламенту; коммерческие базы в Git не попадают.",
    )
    add_checkbox(
        document,
        "Пилот — на ограниченном наборе поставщиков и точек, не на всей сети сразу.",
    )

    remarks_title = document.add_paragraph()
    set_paragraph_spacing(remarks_title, before=10, after=4)
    add_text(remarks_title, "Замечания", size_pt=12.5, bold=True, color=GREEN)
    for _ in range(3):
        line = document.add_paragraph()
        set_paragraph_spacing(line, after=10, before=4)
        add_paragraph_border(line, color="B0B0B0", size="6")
        add_text(line, " ", size_pt=11)

    add_signature_table(document)

    closing = document.add_paragraph()
    set_paragraph_spacing(closing, before=12, after=0)
    add_text(
        closing,
        "После подписи этот лист становится входом ворот G0. Детали реализации не заменяют его "
        "и не расширяют объём без отдельного решения.",
        size_pt=10,
        italic=True,
        color=GRAY,
    )

    core = document.core_properties
    core.author = "ООО «ФИТЭРА»"
    core.last_modified_by = "ООО «ФИТЭРА»"
    core.title = "SNAX — программный контракт для руководителя"
    core.subject = "Краткая форма на подтверждение ворот G0"
    core.category = "Программный контракт"
    core.comments = (
        "Редакция 3.0-exec. Оформление по эталону документов ООО «ФИТЭРА». "
        "Коммерческие выгрузки баз в файл не входят."
    )
    core.created = datetime(2026, 8, 25, 9, 0, tzinfo=UTC)
    core.modified = datetime(2026, 8, 25, 9, 0, tzinfo=UTC)
    core.revision = 1
    return document


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Собрать Word программного контракта в стиле ФИТЭРА."
    )
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        default=DEFAULT_OUTPUT,
        help="Путь к DOCX",
    )
    args = parser.parse_args()
    output: Path = args.output
    output.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(output)
    print(f"wrote {output} ({output.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
