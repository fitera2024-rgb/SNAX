"""Собрать Word технического плана действий в стиле документов ООО «ФИТЭРА»."""

# ruff: noqa: E501

from __future__ import annotations

import argparse
import sys
from datetime import UTC, datetime
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT / "scripts") not in sys.path:
    sys.path.insert(0, str(ROOT / "scripts"))

from fitera_docx import (  # noqa: E402
    CONTENT_WIDTH_MM,
    FILL_YELLOW,
    GRAY,
    GREEN,
    YELLOW_BORDER,
    YELLOW_TEXT,
    add_architecture_cards,
    add_body,
    add_bullet,
    add_callout,
    add_checkbox,
    add_code_block,
    add_data_table,
    add_heading,
    add_kpi_tiles,
    add_letterhead,
    add_numbered,
    add_paragraph_border,
    add_signature_table,
    add_text,
    configure_styles,
    set_paragraph_spacing,
    setup_section,
)

DEFAULT_OUTPUT = ROOT / "docs" / "ФИТЭРА_SNAX_Технический_план_действий_v3.0.docx"


def _title_block(document: Document) -> None:
    kicker = document.add_paragraph()
    set_paragraph_spacing(kicker, after=2, line=1.0, keep_with_next=True)
    add_text(
        kicker,
        "SNAX  ·  ПОДРОБНЫЙ ТЕХНИЧЕСКИЙ ПЛАН ДЕЙСТВИЙ",
        size_pt=10,
        bold=True,
        color=GREEN,
    )
    title = document.add_paragraph("Технический план действий", style="Title")
    title.paragraph_format.keep_with_next = True
    subtitle = document.add_paragraph(
        "Архитектура целевого контура и доработки существующего ландшафта 1С · "
        "очередь работ для архитектора, 1С, backend, QA и Codex",
        style="Subtitle",
    )
    set_paragraph_spacing(subtitle, after=4, line=1.1)
    version = document.add_paragraph()
    set_paragraph_spacing(version, after=10)
    add_text(
        version,
        "Версия 3.0-plan  ·  25 августа 2026 года  ·  Статус: рабочий план для локальной разработки до G0",
        size_pt=10,
        color=GRAY,
    )


def build_document() -> Document:
    document = Document()
    configure_styles(document)
    setup_section(
        document,
        running_header="ФИТЭРА  ·  SNAX  ·  ТЕХНИЧЕСКИЙ ПЛАН ДЕЙСТВИЙ",
        footer_line="ООО «ФИТЭРА»  ·  редакция 3.0-plan  ·  25.08.2026  ·  архитектура и доработки AS-IS",
    )
    add_letterhead(document)
    _title_block(document)

    add_callout(
        document,
        "Назначение документа",
        "Ответить на вопрос «что делать дальше руками»: какие работы по архитектуре "
        "(сервис, контракты, расширение УТ, топология «Форус», выгрузки) и какие доработки "
        "существующего ландшафта 1С (14 расширений, обмены, НСИ, формулы, Frontol, факт продажи) "
        "идут в каком порядке. Календарь — SCHEDULE.md. Полные таблицы действий — "
        "docs/TECHNICAL_ACTION_PLAN.md.",
    )
    add_callout(
        document,
        "Локальный режим",
        "GitHub Actions на дату плана не стартует из‑за billing (D-29). Это не дефект кода. "
        "Quality gate — локальные команды. Не ждать зелёного CI, чтобы мержить WORK-005…009 "
        "по одному в main. Не мержить stacked PR пачкой. Не обещать ускорение G4 из‑за WORK-001…004.",
        fill=FILL_YELLOW,
        border=YELLOW_BORDER,
        title_color=YELLOW_TEXT,
        left_only=True,
    )

    add_kpi_tiles(
        document,
        [
            ("9", "пакетов AP-0…8"),
            ("2", "оси работ"),
            ("G0", "11.09.2026"),
            ("G7", "14.05.2027"),
            ("#10", "первый merge XLSX"),
            ("14", "расширений УТ"),
            ("D-47", "выгрузки"),
            ("D-29", "CI billing"),
        ],
    )

    add_heading(document, "0. Чем план является и чем нет")
    add_body(
        document,
        "План — очередь работ с идентификаторами AP-NNN, владельцами, входом, выходом, "
        "запретом и acceptance. Он не является новым ADR, не утверждает D-01…D-48, "
        "не разрешает payload выгрузок в Git и не разрешает перенос «ПомощникЗакупок» as-is. "
        "Не строим второй расчёт заказа вне 1С.",
    )
    add_data_table(
        document,
        ["Документ", "Роль"],
        [
            ["TECHNICAL_ACTION_PLAN.md", "Полные таблицы AP-001…AP-807"],
            ["SCHEDULE.md", "T0 01.09.2026 → G7 14.05.2027, dumpLagWeeks"],
            ["SPEC.md / TZ addendum", "Как устроены импорт и приёмка"],
            ["ADR-001…004", "Что уже решено"],
            ["IMPLEMENTATION_BACKLOG.md", "Карточки TASK-000…041"],
        ],
        [52, CONTENT_WIDTH_MM - 52],
    )

    add_heading(document, "1. Точка старта, 25.08.2026")
    add_body(
        document,
        "В main: WORK-001…004 (bootstrap, MinIO, outbox/Celery, raw workbook protocol). "
        "Production readers нет. Каталога onec-extension/ нет. UI — тонкая оболочка. "
        "JSON Schema выгрузок уже в contracts/, runtime приёма дампа — нет.",
    )
    add_data_table(
        document,
        ["PR / ветка", "WORK", "TASK", "База"],
        [
            ["#10 work/005-xlsx-reader", "WORK-005", "TASK-007 XLSX", "main"],
            ["work/006-xls-reader", "WORK-006", "TASK-009 isolated XLS", "после #10"],
            ["#12 work/007-csv-reader", "WORK-007", "TASK-008 CSV", "ветка XLS"],
            ["#13 profile schema", "WORK-008", "TASK-010", "ветка CSV"],
            ["#14 profile detection", "WORK-009", "TASK-011", "ветка schema"],
        ],
        [52, 32, 48, CONTENT_WIDTH_MM - 132],
    )
    add_body(
        document,
        "После каждого merge следующую ветку перебазировать на актуальный main и прогнать локальный gate. "
        "Не merge #12/#13/#14, пока предыдущее звено не в main (D-30).",
    )

    add_heading(document, "2. Две оси — не смешивать в одном PR")
    add_architecture_cards(
        document,
        cards=(
            (
                "D9EAF7",
                "Ось A. Архитектура",
                "Сервис, контракты, API, каталоги после дампа, новое расширение УТ. "
                "Спутник только центральной УТ (ADR-004).",
            ),
            (
                "EAF4E3",
                "Ось B. Доработки AS-IS",
                "14 расширений, обмены, MDM магазин–склад–касса, формулы, Frontol, факт продажи D-44. "
                "ПомощникЗакупок разложить, не переносить.",
            ),
            (
                "FFF2CC",
                "Общий запрет",
                "Нет прямого доступа к БД 1С. Непроведённый документ ≠ успех обмена. "
                "Fuzzy не утверждает связь. OCR ≠ факт приёмки.",
            ),
        ),
        note=(
            "Ось A не ждёт полной оси B, кроме пилотов: readers можно делать на синтетике. "
            "Ось B не ждёт готового сервиса: контрольный день можно начать на копии 1С. "
            "Пилот заказа (G4) требует обеих осей. Одна работа — один scope (GOV-006)."
        ),
    )

    add_heading(document, "3. Карта пакетов AP-0…AP-8")
    add_data_table(
        document,
        ["Пакет", "Ось", "Ворота", "Содержание"],
        [
            [
                "AP-0 Готовность",
                "A+B",
                "G0 11.09.2026",
                "Kick-off, активы, merge chain, приём дампов",
            ],
            [
                "AP-1 Сервис без дампов",
                "A",
                "G3 22.01.2027",
                "TASK-012…015, 017…019, каркас TASK-037",
            ],
            ["AP-2 Выгрузки и каталоги", "A", "после D-47 → G1", "TASK-037…040, тестовая копия"],
            ["AP-3 Доработки AS-IS", "B", "G1 / G2", "Обмены, MDM, 14 расширений, формулы, D-44"],
            ["AP-4 Расширение УТ", "A", "после D-01", "onec-extension/, TASK-020…022"],
            ["AP-5 Пилот заказа", "A+B", "G4 19.02.2027", "TASK-016, 023; 3–5 поставщиков"],
            ["AP-6 Пилот приёмки", "A+B", "G5 02.04.2027", "TASK-026…035; D-11…D-14"],
            ["AP-7 Финансы / BI", "B", "G6 23.04.2027", "FIN/SAL/OPS на сверенных источниках"],
            ["AP-8 Rollout", "A+B", "G7 14.05.2027", "TASK-024…025, волны, handover"],
        ],
        [40, 18, 38, CONTENT_WIDTH_MM - 96],
    )
    add_code_block(
        document,
        "AP-0 → AP-2 (дампы) → G1\n"
        "         ↘ AP-3 (формулы/НСИ) → G2 → AP-4 staging → AP-5 G4 → AP-6 G5 → AP-7 G6 → AP-8 G7\n"
        "AP-1 (сервис на синтетике) ──────────────────────────→ G3",
    )
    add_body(
        document,
        "Сдвиг выгрузок — dumpLagWeeks: G1 не раньше чем манифест + 4 недели. "
        "Не догонять сокращением UAT.",
    )

    add_heading(document, "4. Пакет AP-0 — локальная готовность")
    add_body(
        document,
        "К 11.09.2026 программа работает без GitHub Actions, с понятной merge-цепью и готовым приёмом выгрузок. "
        "Не входит: TASK-012+, разбор production DT, написание BSL.",
    )
    add_data_table(
        document,
        ["ID", "Действие", "Не делать", "Acceptance"],
        [
            [
                "AP-001",
                "Kick-off 01.09, freeze инициатив",
                "Стартовать замену 1С/кассы/BI",
                "GOV-001: роли названы",
            ],
            [
                "AP-002",
                "Акт активов прежнего подрядчика",
                "«Потом найдём исходники»",
                "14 расширений или gap в log",
            ],
            [
                "AP-003",
                "Карта баз UNVERIFIED",
                "Кодировать релиз УТ",
                "Код, владелец, гипотеза контура",
            ],
            [
                "AP-004",
                "Черновик манифеста D-47",
                "Класть DT/CF в Git",
                "gitPolicy=DO_NOT_COMMIT_PAYLOAD",
            ],
            ["AP-005", "Merge WORK-005 PR #10", "Ждать CI; merge пачки", "Локальный gate зелёный"],
            [
                "AP-006…009",
                "XLS → CSV → schema → detection",
                "Merge #12 раньше XLS",
                "Каждое звено на свежем main",
            ],
            [
                "AP-010",
                "Зафиксировать локальный gate",
                "Отключать lint из‑за CI",
                "Команды §11 на чистом clone",
            ],
            [
                "AP-011",
                "Эскалация D-29 billing",
                "Чинить workflow кодом",
                "Блокер = биллинг, не дефект",
            ],
            [
                "AP-012",
                "Quarantine bucket дампов",
                "Object key из пути пользователя",
                "Повтор digest = один blob",
            ],
            [
                "AP-013",
                "Steering G0",
                "Утверждать D-20 «на глаз»",
                "Чек-лист G0 или явный перечень дыр",
            ],
        ],
        [24, 48, 52, CONTENT_WIDTH_MM - 124],
    )
    add_body(
        document,
        "До T0 (сейчас → 31.08): не ждать CI; готовить review PR #10; собирать опись активов и черновик D-47; "
        "не начинать TASK-012 в той же ветке, что контракты программы.",
    )

    add_heading(document, "5. Пакет AP-1 — архитектура сервиса без дампов")
    add_body(
        document,
        "Контур «файл → raw → classify → normalize → DQ → import-package → polling API» на синтетических fixtures. "
        "Поглощает ожидание дампов, но не сдвигает G4 влево. Пилотные профили реальных поставщиков — AP-501 после D-02.",
    )
    add_data_table(
        document,
        ["ID", "Результат", "TASK", "Не делать"],
        [
            [
                "AP-101",
                "Row classification, count = raw rows",
                "TASK-012",
                "Молча отбрасывать строку",
            ],
            ["AP-102", "Normalize: GTIN-строка, Decimal", "TASK-013", "Float; coerce в 0/пусто"],
            ["AP-103", "DQ + blocking policy", "TASK-014", "Critical publish"],
            [
                "AP-104",
                "import-package, chunks, digest",
                "TASK-015",
                "Менять raw; повтор без нового run",
            ],
            ["AP-105", "Polling/ACK API", "TASK-018", "Прямой JDBC/COM к БД 1С"],
            ["AP-106", "Mapping cache, 1С побеждает", "TASK-019", "Сервис как master связей"],
            ["AP-107", "Operator UI", "TASK-017", "Дублировать server-side валидацию"],
            ["AP-108", "Каркас registerDumpManifest", "TASK-037", "Загрузка DT в PostgreSQL"],
        ],
        [22, 58, 28, CONTENT_WIDTH_MM - 108],
    )
    add_bullet(
        document,
        "domain без framework/DB; readers возвращают raw, не canonical items; "
        "normalization не открывает файлы; integration_1c только через контракты.",
        bold_prefix="Границы модулей. ",
    )

    add_heading(document, "6. Пакет AP-2 — выгрузки и каталоги")
    add_body(
        document,
        "Старт после D-47. Если манифест позже 11.09.2026 — dumpLagWeeks. "
        "Появление выгрузки наполняет каталоги, а не ломает инварианты.",
    )
    add_data_table(
        document,
        ["ID", "Выход", "Acceptance"],
        [
            ["AP-201", "Intake RECEIVED→REGISTERED", "Повтор SHA = тот же intake"],
            ["AP-202", "Карта баз VERIFIED / gap-list", "Не угадывать релиз УТ (D-01)"],
            ["AP-203", "mdm-object-catalog", "Магазин–склад–касса связаны (MDM-001)"],
            ["AP-204", "extension-passport[]", "ПомощникЗакупок не получает auto KEEP"],
            ["AP-205", "exchange-catalog + store-day", "Непроведённые — issue, не успех"],
            ["AP-206", "Обезличенная тестовая копия", "Не разработка в PROD"],
            ["AP-207", "Compatibility matrix", "CFG-009 на копии, не «на бою»"],
        ],
        [24, 62, CONTENT_WIDTH_MM - 86],
    )

    add_heading(document, "7. Пакет AP-3 — доработки AS-IS 1С")
    add_body(
        document,
        "Стабилизировать текущие базы так, чтобы пилот и P&L не строились на лживых обменах. "
        "Это не перенос ПомощникЗакупок и не замена типовой УТ. Можно параллельно AP-1.",
    )
    add_heading(document, "7.1. Четырнадцать расширений", level=2)
    add_numbered(document, 1, "AP-301 — закрыть GOV-003 / CFG-004: исходники или решение replace.")
    add_numbered(
        document, 2, "AP-302 — паспорт каждого расширения, включая неактивное «ХЛ_PNL» (CFG-008)."
    )
    add_numbered(
        document,
        3,
        "AP-303 — disposition keep / refactor / replace / retire (D-37); Codex не ставит KEEP.",
    )
    add_numbered(
        document, 4, "AP-304 — инвентаризация собственных регистров до отключения (CFG-003)."
    )
    add_numbered(
        document,
        5,
        "AP-305 — разложить «ПомощникЗакупок»: функции → типовой УТ / новое SNAX-расширение / retire (CFG-005, R-26).",
    )
    add_numbered(
        document,
        6,
        "AP-306 — связать «СФ_ДоработкиОбменаСОфлайнОборудованием» с планом Frontol (CFG-006, D-34).",
    )

    add_heading(document, "7.2. Обмены и контрольный день → G1", level=2)
    add_data_table(
        document,
        ["ID", "Действие", "Не делать"],
        [
            ["AP-311", "Трассировка магазин/день INT-001", "Непроведённые «почти ок»"],
            ["AP-312", "Реестр непрошедших обмен FIN-003", "Техзакрытие отрицательных остатков"],
            ["AP-313", "Сверка counts/sums/taxes", "Float в JSON сумм"],
            ["AP-314", "Мониторинг ARC-010", "Ручной «Найти/Выгрузить» как штат"],
            ["AP-315", "Аналитики магазин/склад в регистрах", "Перепроведение сразу в PROD"],
            ["AP-316", "Кассы по точкам с даты второй точки", "Смешивать кассы для отчёта"],
        ],
        [24, 78, CONTENT_WIDTH_MM - 102],
    )

    add_heading(document, "7.3. НСИ, формулы, факт продажи → G2", level=2)
    add_data_table(
        document,
        ["ID", "Действие", "Не делать"],
        [
            ["AP-321", "Таблица магазин–склад–касса–организация", "У каждого свой Excel"],
            ["AP-322", "Синхронизация поставщиков/договоров", "Игнорировать «не удалось найти»"],
            ["AP-331", "Паспорт формулы заказа D-20", "Считать qty во внешнем сервисе"],
            ["AP-332", "Факт продажи D-44", "Выручка = любой проведённый документ"],
            ["AP-333", "Источник себестоимости FIN-006", "Второй движок себестоимости в BI"],
            ["AP-334", "БДР vs ДДС; убрать hardcoded БДС", "Оставить текстовые условия навсегда"],
            ["AP-335", "Проверить данные ХЛ_PNL", "Копировать неподтверждённую методику"],
            ["AP-337", "Реестр отрицательных остатков", "Массовое оприходование «чтобы сошлось»"],
        ],
        [24, 78, CONTENT_WIDTH_MM - 102],
    )
    add_body(
        document,
        "G2 (06.11.2026 в сценарии A) — подписанные паспорта формул, не внедрённый BI. "
        "Production-дашборды до контрольного дня запрещены (GOV-010).",
    )

    add_heading(document, "8. Пакет AP-4 — новое расширение центральной УТ")
    add_body(
        document,
        "Код в onec-extension/. Типовая конфигурация не меняется. Старт после D-01 и TASK-018. "
        "ACK и mapping delta — только после commit транзакции 1С.",
    )
    add_data_table(
        document,
        ["ID", "Результат", "TASK"],
        [
            ["AP-401", "Каркас расширения, права, README", "после D-01"],
            ["AP-402", "Staging: schema, digest, идемпотентность", "TASK-020"],
            ["AP-403", "Matching workspace; fuzzy не approve", "TASK-021"],
            ["AP-404", "Apply цен/упаковок; delta после commit", "TASK-022"],
            ["AP-405", "Auth сервиса ↔ расширения (D-09)", "REQ-NFR-007"],
            ["AP-406", "Запрет проводить поступление из волны заказа", "инварианты 12–13"],
        ],
        [24, 92, CONTENT_WIDTH_MM - 116],
    )

    add_heading(document, "9. Пакеты AP-5 и AP-6 — пилоты")
    add_body(
        document,
        "G4 не раньше 19.02.2027 из‑за WORK-001…004. Старт пилота заказа: G2 + G3 + staging. "
        "OCR (TASK-036) не в волне приёмки. Маршруты DIRECT_TO_STORE / CENTRAL_WAREHOUSE / "
        "RECEIPT_AND_TRANSFER — только после D-11/D-12.",
    )
    add_data_table(
        document,
        ["ID", "Пилот заказа", "ID", "Пилот приёмки"],
        [
            [
                "AP-501",
                "3–5 профилей, golden; не прайсы в Git",
                "AP-601",
                "Receipt domain; OCR ≠ факт",
            ],
            [
                "AP-502",
                "Типовой расчёт УТ, 30–50 SKU",
                "AP-606",
                "РМ магазина; unknown = исключение",
            ],
            [
                "AP-503",
                "Один идемпотентный черновик заказа",
                "AP-608",
                "Exception-only, типовой документ",
            ],
            ["AP-505", "UAT G4; не все поставщики сразу", "AP-610", "32 строки; AC-RCV"],
        ],
        [22, 66, 22, CONTENT_WIDTH_MM - 110],
    )

    add_heading(document, "10. Пакеты AP-7 и AP-8 — контур управления и передача")
    add_bullet(
        document,
        "P&L начисления после D-44 и источника себестоимости; drill-down до документа.",
        bold_prefix="AP-701. ",
    )
    add_bullet(
        document,
        "KPI-passport: сервис не вычисляет показатель; DRAFT ≠ официальный срез.",
        bold_prefix="AP-704. ",
    )
    add_bullet(
        document,
        "Неполный региональный срез нельзя выдавать за всю сеть (SCL-003).",
        bold_prefix="AP-705. ",
    )
    add_bullet(
        document,
        "BI только читает; не исправляет НСИ (BI-001). Казначейство — отдельное D-24.",
        bold_prefix="AP-707/708. ",
    )
    add_bullet(
        document,
        "Restore PostgreSQL + object storage + CF 1С продемонстрирован.",
        bold_prefix="AP-801. ",
    )
    add_bullet(
        document,
        "Новый магазин — чек-лист SCL-004/FIN-018; данные из центральной УТ, не сырой файл.",
        bold_prefix="AP-803/804. ",
    )
    add_bullet(
        document,
        "Frontol cutover только после rollback-теста (D-34). Retire расширения — rehearsal CFG-010.",
        bold_prefix="AP-805/806. ",
    )
    add_body(
        document,
        "Передача: акт G7 14.05.2027. Если D-28 хуже паспорта — сценарий C, G7 около конца июня 2027, не «успеем теми же людьми».",
    )

    add_heading(document, "11. Локальный quality gate")
    add_body(
        document,
        "Пока D-29 открыт, источник истины — локальный прогон. Перед каждым merge в main:",
    )
    add_code_block(
        document,
        "ruff check .\n"
        "ruff format --check .\n"
        "python scripts/validate_manifest.py\n"
        'pytest -q -m "not integration"',
    )
    add_bullet(document, "Не отключать lint/тест «пока CI красный».")
    add_bullet(document, "Не выдавать локальный прогон за прошедший GitHub Actions.")
    add_bullet(document, "Word этого плана: python scripts/build_fitera_tech_plan_docx.py")
    add_bullet(document, "Golden snapshot не обновлять автоматически.")

    add_heading(document, "12. Первые 10 рабочих дней с 01.09.2026")
    add_data_table(
        document,
        ["День", "Дата", "AP", "Результат"],
        [
            ["1", "01.09", "AP-001", "Kick-off, freeze"],
            ["2", "02.09", "AP-002, AP-004", "Акт активов, черновик D-47"],
            ["3", "03.09", "AP-003", "Карта баз UNVERIFIED"],
            ["4–5", "04–07.09", "AP-005…009", "Merge WORK-005…009 по одному"],
            ["6", "08.09", "черновики D-xx", "D-02/D-03/D-20/D-21/D-44 в log, не в коде"],
            ["7", "09.09", "AP-001", "RACI, cadence"],
            ["8", "10.09", "AP-012, AP-108", "Bucket + schema intake"],
            ["9–10", "11.09", "AP-013", "Steering G0"],
        ],
        [18, 24, 36, CONTENT_WIDTH_MM - 78],
    )

    add_heading(document, "13. Что делать разработчику сегодня, до T0")
    add_numbered(document, 1, "Не смешивать docs/контракты с веткой WORK-005.")
    add_numbered(document, 2, "Review и локальный прогон PR #10 (XLSX); merge в main на неделе T0.")
    add_numbered(
        document,
        3,
        "После #10 — retarget XLS → CSV → profile → detection, каждый раз локальный gate.",
    )
    add_numbered(
        document, 4, "Следующая кодовая задача сервиса: TASK-012 (AP-101) отдельной веткой."
    )
    add_numbered(document, 5, "Не создавать onec-extension/ до D-01 как имитацию готовности BSL.")
    add_numbered(
        document, 6, "Не угадывать D-01…D-48: пустые каталоги и synthetic fixtures — можно."
    )

    add_heading(document, "14. Решения, которые план не угадывает")
    add_body(
        document,
        "D-01 релиз УТ · D-02 пилотные поставщики · D-03 упаковки · D-04 цена/пороги · "
        "D-05 частичное apply · D-06 автосопоставление · D-07 номенклатура · D-08 ПДн · "
        "D-09 auth · D-10 тестовая копия · D-11…D-14 приёмка · D-20 формула заказа · "
        "D-24 казначейство · D-25 BI · D-28 ресурсы · D-29 CI billing · D-30 merge chain · "
        "D-31…D-40 топология и расширения · D-44 факт продажи · D-47 выгрузки · D-48 календарь.",
    )

    add_heading(document, "15. Трассировка и риски вне кода")
    add_data_table(
        document,
        ["Пакет", "Основные требования"],
        [
            ["AP-0", "GOV-001…012, D-29, D-30, D-47, D-48"],
            ["AP-1", "IMP-008…014, INT-007, INT-009, REQ-NFR-001…005"],
            ["AP-2", "MDM-001…003, CFG-001, ARC-001…004"],
            ["AP-3", "INT-001…004, ARC-010, CFG-001…010, FIN-001…008, OPS-002"],
            ["AP-4", "INT-006…008, MDM-003, ORD-002"],
            ["AP-5 / AP-6", "ORD-001…012, RCV-001…014"],
            ["AP-7 / AP-8", "FIN-009…018, BI-001, SCL-004, CFG-010"],
        ],
        [32, CONTENT_WIDTH_MM - 32],
    )
    add_body(
        document,
        "Риски, которые план не лечит кодом: D-29 — эскалация оплаты; нет дампов — dumpLagWeeks; "
        "D-28 ниже паспорта — сценарий C; закрытый код — CFG-004; ПомощникЗакупок as-is — запрет AP-305.",
    )

    document.add_page_break()
    add_heading(document, "16. Лист согласования")
    add_body(document, "Отметьте итог:", after=4)
    add_checkbox(
        document, " план действий v3.0-plan как рабочую очередь до G0.", bold_lead="Утверждаю"
    )
    add_checkbox(document, " с замечаниями.", bold_lead="Утверждаю")
    add_checkbox(document, " — вернуть на доработку.", bold_lead="Не утверждаю")
    add_heading(document, "Обязательные инварианты плана", level=2)
    add_checkbox(document, "Не строим второй расчёт заказа вне 1С.")
    add_checkbox(document, "Нет прямого доступа сервиса к БД 1С.")
    add_checkbox(document, "Stacked PR WORK-005…009 мержить по одному, локальный gate вместо CI.")
    add_checkbox(document, "ПомощникЗакупок не переносится as-is.")
    add_checkbox(document, "Payload выгрузок баз в Git не попадает.")
    add_checkbox(document, "G4 не ускоряется из‑за WORK-001…004.")

    remarks = document.add_paragraph()
    set_paragraph_spacing(remarks, before=10, after=4)
    add_text(remarks, "Замечания", size_pt=12.5, bold=True, color=GREEN)
    for _ in range(3):
        line = document.add_paragraph()
        set_paragraph_spacing(line, after=10, before=4)
        add_paragraph_border(line, color="B0B0B0", size="6")
        add_text(line, " ", size_pt=11)

    add_signature_table(
        document,
        [
            "Спонсор / собственник",
            "Владелец 1С",
            "Технический лидер сервиса",
            "Финансовый директор",
            "Руководитель категорийного менеджмента",
            "Руководитель розницы",
            "Операционный директор",
            "Координатор / ФИТЭРА",
        ],
    )
    closing = document.add_paragraph()
    set_paragraph_spacing(closing, before=12, after=0)
    add_text(
        closing,
        "После утверждения этот документ становится операционным входом G0 вместе с календарём v3.0. "
        "Полные строки AP-NNN остаются в TECHNICAL_ACTION_PLAN.md.",
        size_pt=10,
        italic=True,
        color=GRAY,
    )

    core = document.core_properties
    core.author = "ООО «ФИТЭРА»"
    core.last_modified_by = "ООО «ФИТЭРА»"
    core.title = "SNAX — технический план действий"
    core.subject = "Архитектура и доработки AS-IS v3.0-plan"
    core.category = "Технический план"
    core.comments = (
        "Редакция 3.0-plan. Очередь работ по архитектуре сервиса и доработкам ландшафта 1С. "
        "Коммерческие выгрузки баз в файл не входят."
    )
    core.created = datetime(2026, 8, 25, 12, 0, tzinfo=UTC)
    core.modified = datetime(2026, 8, 25, 12, 0, tzinfo=UTC)
    core.revision = 1
    return document


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Собрать Word технического плана действий в стиле ФИТЭРА."
    )
    parser.add_argument("-o", "--output", type=Path, default=DEFAULT_OUTPUT, help="Путь к DOCX")
    args = parser.parse_args()
    output: Path = args.output
    output.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(output)
    print(f"wrote {output} ({output.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
